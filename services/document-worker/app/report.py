"""Report generation: compliance reports, summaries and evidence matrices.

Every row in a generated report comes from a verified citation supplied by the API. This
module formats; it never asserts. A row whose citation failed verification arrives with
`verified: false` and is rendered as such rather than being dropped, so the reader sees the
gap instead of a silently shorter table.
"""

from __future__ import annotations

import csv
import io
from typing import Any

from .pdfkit import (
    BOLD,
    DANGER,
    Document,
    INK,
    MUTED,
    NAVY,
    SUCCESS,
    WARNING,
)

RESULT_COLORS = {
    "compliant": SUCCESS,
    "non_compliant": DANGER,
    "needs_evidence": WARNING,
    "not_assessed": MUTED,
}

RESULT_LABELS = {
    "compliant": "Compliant",
    "non_compliant": "Non-compliant",
    "needs_evidence": "Needs evidence",
    "not_assessed": "Not assessed",
}


def build_pdf(payload: dict[str, Any]) -> bytes:
    """Renders a report PDF.

    Laid out as a document somebody files, not a screen printed to paper: a running head on
    every page, a table at the top saying what this is and what it concluded, numbered
    sections, and findings in a table rather than a scroll of stacked paragraphs.

    Layout goes through `pdfkit.Document`, which measures every line before drawing it, so
    a long evidence matrix paginates rather than silently truncating.
    """
    title = payload.get("title") or "Compliance report"
    doc = Document(
        header_left="UXE Consulting AI - Compliance report",
        header_right=f"Generated {str(payload.get('generatedAt', ''))[:10]}",
        footer_text="UXE Consulting AI  -  verified answers, exact evidence",
    )

    doc.text(title, size=19, color=NAVY, bold=True, gap=4)
    if payload.get("subtitle"):
        doc.text(
            f"Compliance review of the submitted documents against the approved knowledge base.  -  {payload['subtitle']}",
            size=9,
            color=MUTED,
            gap=12,
        )

    # --- What this document is, before anything it concludes ------------------
    decision = payload.get("decision")
    qualifier = payload.get("decisionQualifier")
    verdict = (qualifier or decision or "Not determined").upper()
    coverage = float(payload.get("coverage") or 0) * 100
    confidence = float(payload.get("confidence") or 0) * 100
    rows = payload.get("rows") or []
    counts = _count_results(rows)

    doc.meta_table(
        [
            ("Verdict", verdict),
            ("Reviewed", _documents_line(payload.get("documentsReviewed") or [])),
            (
                "Result",
                f"{counts['compliant']} met, {counts['non_compliant']} not met, "
                f"{counts['needs_evidence']} cannot be verified from the documents supplied."
                if rows
                else "The evidence matrix is not included in this edition.",
            ),
            ("Confidence", f"Evidence coverage {coverage:.0f}%  -  Confidence {confidence:.0f}%"),
        ]
    )

    section = 0

    section += 1
    doc.section(section, "Summary")
    doc.text(payload.get("summary") or "", size=9.5, gap=14)

    documents = payload.get("documentsReviewed") or []
    if documents:
        section += 1
        doc.section(section, "Documents reviewed")
        doc.table(
            ["Document", "Version", "Role", "Pages"],
            [
                [
                    str(item.get("title", "")),
                    str(item.get("version", "")),
                    "Knowledge base" if item.get("role") == "governing" else "Submitted",
                    str(item.get("pages") or "-"),
                ]
                for item in documents
            ],
            widths=[0.52, 0.12, 0.24, 0.12],
        )

    assumptions = payload.get("assumptions") or []
    if assumptions:
        section += 1
        doc.section(section, "Scope of this review")
        for item in assumptions:
            doc.text(f"-  {item}", size=9, color=MUTED, gap=4)
        doc.space(8)

    if rows:
        section += 1
        doc.section(section, "Findings")
        doc.text(
            "One row per requirement tested. The clause and page are where the requirement "
            "comes from; the quoted text is what the submitted document says.",
            size=9,
            color=MUTED,
            gap=10,
        )
        for index, row in enumerate(rows, start=1):
            _draw_finding(doc, index, row)

    recommendations = payload.get("recommendations") or []
    if recommendations:
        section += 1
        doc.section(section, "What to do next")
        doc.table(
            ["Priority", "Action"],
            [
                [str(item.get("priority", "")).upper(), str(item.get("action", ""))]
                for item in recommendations
            ],
            widths=[0.16, 0.84],
            emphasis=[0],
        )

    disclosures = payload.get("disclosures") or []
    if disclosures:
        doc.callout("How to read this.", "  ".join(str(item) for item in disclosures))

    return doc.finish()


def _has_evidence(rows: list[dict[str, Any]]) -> bool:
    """Whether any row carries a citation.

    When evidence is excluded the API sends the rows with their source, locator, page and
    excerpt blanked. Rendering those as empty columns would leave four dead columns in a
    spreadsheet and an empty quotation mark in a document, so each builder drops them
    instead of printing nothing into them.
    """
    return any((row.get("source") or "").strip() for row in rows)


def _count_results(rows: list[dict[str, Any]]) -> dict[str, int]:
    counts = {"compliant": 0, "non_compliant": 0, "needs_evidence": 0, "not_assessed": 0}
    for row in rows:
        key = row.get("result", "not_assessed")
        if key in counts:
            counts[key] += 1
    return counts


def _documents_line(documents: list[dict[str, Any]]) -> str:
    knowledge = [d.get("title", "") for d in documents if d.get("role") == "governing"]
    submitted = [d.get("title", "") for d in documents if d.get("role") != "governing"]
    parts = []
    if submitted:
        parts.append(f"Submitted: {', '.join(submitted)}")
    if knowledge:
        parts.append(f"Against: {', '.join(knowledge)}")
    return "  -  ".join(parts) or "-"


def _draw_finding(doc: Document, index: int, row: dict[str, Any]) -> None:
    """One requirement, its verdict, and — when present — where the evidence came from.

    The source line and the quotation are drawn only if the payload carries them. That is
    what "include evidence" turns off: the requirement and the verdict stay, the citation
    block starting with the document name goes.
    """
    result = row.get("result", "not_assessed")
    requirement = f"{index}.  {row.get('requirement', '')}"
    excerpt = (row.get("excerpt") or "").strip()
    source = str(row.get("source") or "").strip()

    doc.keep_together(
        doc.measure(requirement, 10.5, bold=True)
        + doc.measure(str(row.get("finding", "")), 9.5)
        + (doc.measure(excerpt, 9, indent=14) if excerpt else 0)
        + 52
    )
    doc.text(requirement, size=10.5, bold=True, gap=3)
    doc.text(
        RESULT_LABELS.get(result, result).upper(),
        size=8,
        color=RESULT_COLORS.get(result, MUTED),
        bold=True,
        gap=4,
    )
    doc.text(str(row.get("finding", "")), size=9.5, gap=5)

    if source:
        page = f"  -  p. {row.get('page')}" if row.get("page") else ""
        verified = "verified" if row.get("verified") else "UNVERIFIED"
        doc.text(
            f"{source} ({row.get('version', '-')})  -  {row.get('location') or '-'}{page}  -  [{verified}]",
            size=8,
            color=MUTED if row.get("verified") else WARNING,
            gap=4,
        )
    if excerpt:
        doc.text(f'"{excerpt}"', size=9, color=MUTED, indent=14, gap=8)
    doc.rule(gap=10)


def build_docx(payload: dict[str, Any]) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor

    document = docx.Document()

    brand = document.add_paragraph()
    brand_run = brand.add_run("UXE Consulting AI")
    brand_run.bold = True
    brand_run.font.color.rgb = RGBColor(0x31, 0x56, 0xF5)
    brand_run.font.size = Pt(11)

    document.add_heading(payload.get("title") or "Compliance report", level=0)
    if payload.get("subtitle"):
        document.add_paragraph(payload["subtitle"])
    document.add_paragraph(f"Generated {payload.get('generatedAt', '')}")

    decision = payload.get("decision")
    if decision:
        paragraph = document.add_paragraph()
        run = paragraph.add_run((payload.get("decisionQualifier") or decision).upper())
        run.bold = True
        run.font.color.rgb = (
            RGBColor(0xE5, 0x48, 0x4D) if decision == "no" else RGBColor(0x12, 0xA8, 0x6B)
        )

    document.add_paragraph(
        f"Evidence coverage {float(payload.get('coverage') or 0) * 100:.0f}%  |  "
        f"Confidence {float(payload.get('confidence') or 0) * 100:.0f}%"
    )

    document.add_heading("Summary", level=1)
    document.add_paragraph(payload.get("summary") or "")

    documents = payload.get("documentsReviewed") or []
    if documents:
        document.add_heading("Documents reviewed", level=1)
        for doc in documents:
            document.add_paragraph(
                f"{doc.get('title')} ({doc.get('version')}, {doc.get('role')})", style="List Bullet"
            )

    assumptions = payload.get("assumptions") or []
    if assumptions:
        document.add_heading("Assumptions and scope limits", level=1)
        for item in assumptions:
            document.add_paragraph(item, style="List Bullet")

    rows = payload.get("rows") or []
    if rows:
        with_evidence = _has_evidence(rows)
        document.add_heading("Findings" if not with_evidence else "Evidence matrix", level=1)
        headers = ["Requirement", "Result", "Finding"]
        if with_evidence:
            headers += ["Source", "Location", "Excerpt"]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("requirement", ""))
            cells[1].text = RESULT_LABELS.get(row.get("result", ""), str(row.get("result", "")))
            cells[2].text = str(row.get("finding", ""))
            if with_evidence:
                cells[3].text = f"{row.get('source', '')} ({row.get('version', '')})"
                page = f" p.{row.get('page')}" if row.get("page") else ""
                verified = "" if row.get("verified") else " [UNVERIFIED]"
                cells[4].text = f"{row.get('location', '')}{page}{verified}"
                cells[5].text = str(row.get("excerpt", ""))

    recommendations = payload.get("recommendations") or []
    if recommendations:
        document.add_heading("Recommended actions", level=1)
        for item in recommendations:
            document.add_paragraph(
                f"[{str(item.get('priority', '')).upper()}] {item.get('action', '')}",
                style="List Bullet",
            )

    disclosures = payload.get("disclosures") or []
    if disclosures:
        document.add_heading("Disclosures", level=1)
        for item in disclosures:
            document.add_paragraph(item, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


EVIDENCE_HEADERS = [
    "Requirement",
    "Result",
    "Finding",
    "Source",
    "Version",
    "Chapter / Section / Clause",
    "Page",
    "Exact location",
    "Supporting excerpt",
    "Confidence",
    "Verified",
]


# Columns that only exist when the evidence behind a finding is being reproduced.
_EVIDENCE_COLUMNS = (3, 4, 5, 6, 7, 8, 10)


def _headers(with_evidence: bool) -> list[str]:
    if with_evidence:
        return EVIDENCE_HEADERS
    return [h for i, h in enumerate(EVIDENCE_HEADERS) if i not in _EVIDENCE_COLUMNS]


def _evidence_row(row: dict[str, Any], with_evidence: bool = True) -> list[str]:
    cells = [
        str(row.get("requirement", "")),
        RESULT_LABELS.get(row.get("result", ""), str(row.get("result", ""))),
        str(row.get("finding", "")),
        str(row.get("source", "")),
        str(row.get("version", "")),
        str(row.get("location", "")),
        str(row.get("page") or ""),
        str(row.get("location", "")),
        str(row.get("excerpt", "")),
        f"{float(row.get('confidence') or 0) * 100:.0f}%",
        "yes" if row.get("verified") else "no",
    ]
    if with_evidence:
        return cells
    return [c for i, c in enumerate(cells) if i not in _EVIDENCE_COLUMNS]


def build_csv(payload: dict[str, Any]) -> bytes:
    rows = payload.get("rows") or []
    with_evidence = _has_evidence(rows)
    buffer = io.StringIO()
    writer = csv.writer(buffer, quoting=csv.QUOTE_ALL)
    writer.writerow(_headers(with_evidence))
    for row in rows:
        writer.writerow(_evidence_row(row, with_evidence))
    # BOM so Excel opens UTF-8 correctly on Windows without mangling accented characters.
    return b"\xef\xbb\xbf" + buffer.getvalue().encode("utf-8")


def build_xlsx(payload: dict[str, Any]) -> bytes:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Evidence matrix"

    header_fill = PatternFill("solid", fgColor="3156F5")
    header_font = Font(color="FFFFFF", bold=True)

    rows = payload.get("rows") or []
    with_evidence = _has_evidence(rows)

    sheet.append(_headers(with_evidence))
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(vertical="center", wrap_text=True)

    for row in rows:
        sheet.append(_evidence_row(row, with_evidence))

    all_widths = [34, 16, 58, 28, 12, 26, 8, 26, 64, 12, 10]
    widths = (
        all_widths
        if with_evidence
        else [w for i, w in enumerate(all_widths) if i not in _EVIDENCE_COLUMNS]
    )
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[openpyxl.utils.get_column_letter(index)].width = width

    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    # Freeze the header so a long matrix stays readable while scrolling.
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions

    summary = workbook.create_sheet("Summary")
    summary.append(["Title", payload.get("title", "")])
    summary.append(["Generated", payload.get("generatedAt", "")])
    summary.append(["Decision", payload.get("decisionQualifier") or payload.get("decision") or ""])
    summary.append(["Evidence coverage", f"{float(payload.get('coverage') or 0) * 100:.0f}%"])
    summary.append(["Confidence", f"{float(payload.get('confidence') or 0) * 100:.0f}%"])
    summary.append([])
    summary.append(["Summary", payload.get("summary", "")])
    summary.append([])
    summary.append(["Assumptions"])
    for item in payload.get("assumptions") or []:
        summary.append(["", item])
    summary.append([])
    summary.append(["Disclosures"])
    for item in payload.get("disclosures") or []:
        summary.append(["", item])
    summary.column_dimensions["A"].width = 22
    summary.column_dimensions["B"].width = 100

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def build_markdown(payload: dict[str, Any]) -> bytes:
    lines: list[str] = [
        f"# {payload.get('title', 'Compliance report')}",
        "",
        payload.get("subtitle") or "",
        "",
        f"_Generated {payload.get('generatedAt', '')}_",
        "",
    ]

    if payload.get("decision"):
        lines.append(f"**{(payload.get('decisionQualifier') or payload['decision']).upper()}**")
        lines.append("")

    lines.extend(
        [
            f"Evidence coverage **{float(payload.get('coverage') or 0) * 100:.0f}%** - "
            f"Confidence **{float(payload.get('confidence') or 0) * 100:.0f}%**",
            "",
            "## Summary",
            "",
            payload.get("summary") or "",
            "",
        ]
    )

    if payload.get("documentsReviewed"):
        lines.extend(["## Documents reviewed", ""])
        for doc in payload["documentsReviewed"]:
            lines.append(f"- {doc.get('title')} ({doc.get('version')}, {doc.get('role')})")
        lines.append("")

    if payload.get("assumptions"):
        lines.extend(["## Assumptions and scope limits", ""])
        for item in payload["assumptions"]:
            lines.append(f"- {item}")
        lines.append("")

    markdown_rows = payload.get("rows") or []
    if markdown_rows:
        with_evidence = _has_evidence(markdown_rows)
        headers = ["Requirement", "Result", "Finding"]
        if with_evidence:
            headers += ["Source", "Location", "Page", "Excerpt"]
        headers += ["Confidence", "Verified"] if with_evidence else ["Confidence"]
        lines.extend(
            [
                "## Evidence matrix" if with_evidence else "## Findings",
                "",
                "| " + " | ".join(headers) + " |",
                "| " + " | ".join("---" for _ in headers) + " |",
            ]
        )
        for row in markdown_rows:
            cells = [
                str(row.get("requirement", "")),
                RESULT_LABELS.get(row.get("result", ""), ""),
                str(row.get("finding", "")),
            ]
            if with_evidence:
                cells += [
                    str(row.get("source", "")),
                    str(row.get("location", "")),
                    str(row.get("page") or ""),
                    str(row.get("excerpt", "")),
                ]
            cells.append(f"{float(row.get('confidence') or 0) * 100:.0f}%")
            if with_evidence:
                cells.append("yes" if row.get("verified") else "no")
            # Escape pipes so a citation containing one cannot break the table.
            lines.append("| " + " | ".join(c.replace("|", "\\|").replace("\n", " ") for c in cells) + " |")
        lines.append("")

    if payload.get("recommendations"):
        lines.extend(["## Recommended actions", ""])
        for item in payload["recommendations"]:
            lines.append(f"- **[{str(item.get('priority', '')).upper()}]** {item.get('action', '')}")
        lines.append("")

    if payload.get("disclosures"):
        lines.extend(["## Disclosures", ""])
        for item in payload["disclosures"]:
            lines.append(f"- {item}")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


BUILDERS = {
    "pdf": (build_pdf, "application/pdf", "pdf"),
    "docx": (
        build_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "xlsx": (
        build_xlsx,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "xlsx",
    ),
    "csv": (build_csv, "text/csv", "csv"),
    "markdown": (build_markdown, "text/markdown", "md"),
}
